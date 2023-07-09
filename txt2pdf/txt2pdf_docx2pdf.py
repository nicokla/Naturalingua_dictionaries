import os
import os.path
from os import path
import glob

import sys
sys.path.append('/Users/nicolas/Desktop/NaturaLingua/utils')
from utils import createNecessaryFolders, getListFiles

# ==========================
# https://python-docx.readthedocs.io/en/latest/user/install.html
# https://github.com/AlJohri/docx2pdf

from docx import Document
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx2pdf import convert
from subprocess import Popen, PIPE

def add_hyperlink(paragraph, text, url):
	part = paragraph.part
	r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
	hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
	hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
	new_run = docx.oxml.shared.OxmlElement('w:r')
	rPr = docx.oxml.shared.OxmlElement('w:rPr')
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)
	r = paragraph.add_run ()
	r._r.append (hyperlink)
	r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
	r.font.underline = True
	return hyperlink


def dispatchPdfs(dico):
	for file in getListFiles(['/Users/nicolas/Desktop/NaturaLingua/temp/~*']):
		os.remove(file)
	p_status=1
	while(p_status != 0):
		process = Popen(['docx2pdf', '--keep-active', '/Users/nicolas/Desktop/NaturaLingua/temp'])
		# stdout, stderr = process.communicate()
		p_status = process.wait()
		for weirdPath in dico:
			normalPath=dico[weirdPath]
			weirdPathPdf=f'{weirdPath[:-5]}.pdf'
			normalPathPdf=f'{normalPath[:-5]}.pdf'
			if path.exists(weirdPathPdf):
				createNecessaryFolders(normalPathPdf)
				os.rename(weirdPathPdf, normalPathPdf)
				os.remove(weirdPath)
		for file in getListFiles(['/Users/nicolas/Desktop/NaturaLingua/temp/*.docx']):
			os.remove(file)


# =============================

def getAuthor(title):
	dicoAuthors={
		'Jules et Jim':'François Truffaut',
		'Les 400 coups':'François Truffaut',
		'Les parapluies de Cherbourg':'Jacques Demy',
		'Amelie':'Jean-Pierre Jeunet',
		'Turtles can fly':'Bahman Ghobadi',
		'Leili baa man ast':'Kamal Tabrizi',
		'A separation':'Asghar Farhadi',
		'Very big shot':'Mir-Jean Bou Chaaya',
		'The Kite':'Randa Chahal Sabbag',
		'Heaven without people':'Lucien Bourjeily',
		'Ghadi':'Amin Dora',
		'Beirut Oh Beirut':'Maroun Bagdadi',
		'Shtisel':'Ori Elon and Yehonatan Indursky',
		'Vizontele':'Yılmaz Erdoğan',
		'The butterfly\'s dream':'Yılmaz Erdoğan',
		'Gora':'Ömer Faruk Sorak',
		'Ulysses gaze':'Theo Angelopoulos',
		'Eleni':'Peter Yates',
		'The life ahead':'Edoardo Ponti',
		'Ladri di biciclette':'Vittorio De Sica',
		'La dolce vita':'Federico Fellini',
		'Il sorpasso':'Dino Risi',
		'Udta Punjab':'Abhishek Chaubey',
		'Guilty':'Ruchi Narain',
		'Black Friday':'Anurag Kashyap',
		'The Scent of Green Papaya':'Tran Anh Hung',
		'Furie':'Lê Văn Kiệt',
		'Okja':'Bong Joon-ho',
		'Night in paradise':'Park Hoon-Jung',
		'Descendants of the sun':'Kim Eun-sook and Kim Won-seok',
		'Us and them':'Rene Liu',
		'In the heat of the sun':'Jiang Wen',
		'Big fish and begonia':'Xuan Liang and Chun Zhang',
		'A brighter summer day':'Edward Yang',
		'Totoro':'Hayao Miyazaki',
		'Tokyo godfathers':'Satoshi Kon and Shôgo Furuya',
		'Rashomon':'Akira Kurosawa',
		'Ponyo':'Hayao Miyazaki',
		'Nausicaa':'Hayao Miyazaki',
		'Mononoke':'Hayao Miyazaki',
		'Mirai no Mirai':'Mamoru Hosoda',
		'The Wind Rises':'Hayao Miyazaki',
		'From Up on Poppy Hill':'Gorō Miyazaki',
		'Kiki\'s Delivery Service':'Hayao Miyazaki',
		'Chihiro':'Hayao Miyazaki',
		'Castle in the Sky':'Hayao Miyazaki',
		'The Howl Moving Castle':'Hayao Miyazaki',
		'Leto':'Kirill Serebrennikov',
		'Stalker':'Andreï Tarkovski',
		'The cranes are flying':'Mikhail Kalatozov',
		'Pixote':'Héctor Babenco',
		'Carandiru':'Miguel Gonçalves',
		'Central do Brasil':'Walter Salles',
		'Ilha das flores':'Jorge Furtado',
		'Invisible life':'Eurídice Gusmão',
		'José e Pilar':'Miguel Gonçalves Mendes',
		'Cidade de Deus':'Fernando Meirelles and Kátia Lund',
		'The salt of the earth':'Juliano Ribeiro Salgado and Wim Wenders',
		'Roma':'Alfonso Cuarón',
		'Sin nombre':'Cary Joji Fukunaga',
		'Ya no estoy aqui':'Fernando Frías de la Parra',
		'Girl from nowhere':'Kondej Jaturanrasamee',
		'Monkey twins':'Nonthakorn Thaweesuk'
	}
	if title in dicoAuthors:
		return dicoAuthors[title]
	elif title.startswith('Shtisel'):
		return dicoAuthors['Shtisel']
	elif title.startswith('Descendants of the sun'):
		return dicoAuthors['Descendants of the sun']
	elif title.startswith('Girl from nowhere'):
		return dicoAuthors['Girl from nowhere']
	elif title.startswith('Monkey twins'):
		return dicoAuthors['Monkey twins']


def doAllPdfs_movies():
	# textFiles=getListFiles(['/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/movies/*/**.txt'])
	textFiles=getListFiles(['/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/movies/Chinese/*.txt'])
	dico={}
	problems=[]
	for fileName in textFiles:
		liste=fileName.split('/')
		title=liste[-1][:-4]
		author = getAuthor(title)
		if author == '':
			print('Problem!')
			problems.append(fileName)
		language=liste[-2]
		pathOutputNormal=f'/Users/nicolas/Desktop/NaturaLingua/toEnglish/{language}/movies/{title}.docx'
		pathOutput=f'/Users/nicolas/Desktop/NaturaLingua/temp/{title}.docx'
		print(pathOutput)
		dico[pathOutput]=pathOutputNormal
		document = Document('/Users/nicolas/Desktop/NaturaLingua/txt2pdf/modele.docx')
		file=open(fileName,'r')
		content=file.read()
		file.close()
		for paragraph in document.paragraphs:
			for run in paragraph.runs:
				if('title' in run.text):
					run.text = title
				if('author' in run.text):
					run.text = author
				if('content' in run.text):
					run.text = content
		document.save(pathOutput)
	dispatchPdfs(dico)
	print('problems',problems)


# =================

def doAllPdfs_docs():
	textFiles=getListFiles(['/Users/nicolas/Desktop/NaturaLingua/tatoeba/wordsToPhrasesAlpha/*.txt',
	'/Users/nicolas/Desktop/NaturaLingua/tatoeba/wordsToPhrasesFreq/*.txt'])
	dico={}
	for fileName in textFiles:
		fromEnglish=False
		if('english2' in fileName):
			fromEnglish=True
			index1=fileName.find('english2')+8
		else:
			liste=fileName.split('/')
			name=liste[-1]
			index1=len(fileName)-len(name)
		toto=fileName.find('_small')
		isSmall=(toto!=-1)
		if(not isSmall):
			language=fileName[index1:-4].capitalize()
		else:
			language=fileName[index1:toto].capitalize()
		keyword='frequency'
		if('Alpha' in fileName):
			keyword='alphabetical'
		title=f'{language} sentences (words by {keyword} order)'
		if(isSmall):
			title= title + ', frequent words'
		if fromEnglish:
			title = 'English to ' + title
		keyword2='freq'
		if('Alpha' in fileName):
			keyword2='alpha'
		titleDoc=f'{language} sentences ({keyword2})'
		if(isSmall):
			titleDoc= titleDoc + ', frequent words'
		if fromEnglish:
			titleDoc = 'English to ' + titleDoc
		if(fromEnglish):
			pathOutputNormal=f'/Users/nicolas/Desktop/NaturaLingua/toEnglish/{language}/from english/{titleDoc}.docx'
		else:
			pathOutputNormal=f'/Users/nicolas/Desktop/NaturaLingua/toEnglish/{language}/to english/{titleDoc}.docx'
		pathOutput=f'/Users/nicolas/Desktop/NaturaLingua/temp/{titleDoc}.docx'
		print(pathOutput)
		dico[pathOutput]=pathOutputNormal
		document = Document('/Users/nicolas/Desktop/NaturaLingua/tatoeba/wordsToPhrasesFreq/template.docx')
		file=open(fileName,'r')
		content=file.read()
		file.close()
		for paragraph in document.paragraphs:
			for run in paragraph.runs:
				if('title' in run.text):
					run.text = title
				if('content' in run.text):
					run.text = content
		document.save(pathOutput)
	dispatchPdfs(dico)

# =============================

def doAllPdfs_YoutubeVideos():
	channelsFileNames=getListFiles(['/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/youtube/*/*.txt'])
	for index, channelFileName in enumerate(channelsFileNames):
		liste=channelFileName.split('/')
		language=liste[-2]
		if language.lower() == 'english':
			continue
		channelName=liste[-1][:-4]
		parentNormal=f'/Users/nicolas/Desktop/NaturaLingua/toEnglish/{language}/youtube/{channelName}'
		parentFuckedUp='/Users/nicolas/Desktop/NaturaLingua/temp'
		# os.chmod(parent, stat.S_IRWXU | stat.S_IRWXG | stat.S_IRWXO)
		file=open(channelFileName,'r')
		lines=file.readlines()
		i=1
		numVids=0
		dico={}
		print(language, channelName)
		while i<len(lines):
			while i<len(lines) and not lines[i].startswith('------'):
				i+=1
			if i >= len(lines):
				break
			print(f'\r {numVids} {i/len(lines)*100} %', end='', flush=True)
			if lines[i].startswith('------') and len(lines[i-1]) <= 1:
				title=lines[i+1][:-1]
				link=lines[i+2][:-1]
				i = i+4
				pathOutput=f'{parentFuckedUp}/{link[-11:]}.docx'
				titleVid=title[:240].replace('/','_')
				pathOutputNormal=f'{parentNormal}/{titleVid}.docx'
				if path.exists(f'{pathOutputNormal[:-5]}.pdf'):
					numVids+=1
					continue
			dico[pathOutput] = pathOutputNormal					
			document = Document('/Users/nicolas/Desktop/NaturaLingua/txt2pdf/modele123.docx')
			paragraph=document.add_paragraph()
			paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			aaa=paragraph.add_run(title+'\n')
			aaa=add_hyperlink(paragraph, link+'\n', link)
			paragraph2=document.add_paragraph('')
			paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT
			while(i<len(lines) and not lines[i].startswith('------')):
				aaa=paragraph2.add_run(lines[i])
				i+=1
			createNecessaryFolders(pathOutput)
			document.save(pathOutput)
			numVids+=1
		dispatchPdfs(dico)


	

















# ==========================================================


def txt2word(fileNameTxt):
	fileNameDocx = f'{fileNameTxt[:-4]}.docx'
	document = Document(modeleDocx)
	fileTxt=open(fileNameTxt,'r')
	lines=fileTxt.readlines()
	paragraph=document.add_paragraph('')
	for line in lines:
		if line.startswith('https://youtu.be'):
			add_hyperlink(paragraph, line, line[:-1])
		else:
			paragraph.add_run(line)
	document.save(fileNameDocx)

import os
import os.path
from os import path
import glob
from docx2pdf import convert

def getListFiles(patterns):
	listeFiles=[]
	for aaa in patterns:
		files=glob.glob(aaa)
		listeFiles += files
	for a in listeFiles:
		if a.__contains__('readme'):
			listeFiles.remove(a)
	return listeFiles

def txt2pdfNew():
	parent = '/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/youtube'
	listFileNames=getListFiles([f'{parent}/**/*.txt'])
	for fileName in listFileNames:
		print(fileName)
		txt2word(fileName)
	for path in glob.glob(f'{parent}/*'):
		if(os.path.isdir(path)):
			print(path)
			convert(path, keep_active=True)

# txt2pdfNew()

def txt2pdf(dir):
	obj = os.scandir(dir)
	for entry in obj :
		if entry.is_file() and entry.path.endswith('.txt') and not entry.name.startswith('readme'):
			txt2word(entry.path)
	convert(dir, keep_active=True)



# --------------------

def doSomething_1():
	dirs=['/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/youtube/']
	# '/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/movies/',
	names=['youtube']#'movies',
	dir2='/Users/nicolas/Desktop/NaturaLingua/toEnglish'
	for aaa,dir in enumerate(dirs):
		obj = os.scandir(dir)
		for entry in obj :
			if entry.is_dir():
				# txt2pdf(entry.path)
				if(not path.exists(f'{dir2}/{entry.name}')):
					os.mkdir(f'{dir2}/{entry.name}')
				if(not path.exists(f'{dir2}/{entry.name}/{names[aaa]}')):
					os.mkdir(f'{dir2}/{entry.name}/{names[aaa]}')
				obj2 = os.scandir(entry.path)
				for entry2 in obj2 :
					if entry2.is_file() and entry2.path.endswith('.pdf'):
						os.rename(entry2.path, f'{dir2}/{entry.name}/{names[aaa]}/{entry2.name}')


# ==============================

import tarfile
import os.path
import os

dir='/Users/nicolas/Desktop/NaturaLingua/toEnglish'

def make_tarfile(output_filename, source_dir):
	with tarfile.open(output_filename, "w:gz") as tar:
		tar.add(source_dir, arcname=os.path.basename(source_dir))

def doSomething_3():
	obj = os.scandir(dir)
	for entry in obj :
		if entry.is_dir():
			make_tarfile(f'{entry.path}.tar.gz', entry.path)







# -----------------
# titles

import os.path
import os
from docx import Document
from docx2pdf import convert

def createTitle(title, author, language):
	fileName = '/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/txt2pdf/Title.docx'
	file2Name='/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/txt2pdf/Title_after.docx'
	filePdf = f'{file2Name[:-5]}.pdf'
	document = Document(fileName)
	for paragraph in document.paragraphs:
		for run in paragraph.runs:
			run.text = run.text.replace('title', title)
			run.text = run.text.replace('author', author)
			run.text = run.text.replace('language', language)
	document.save(file2Name)
	convert(file2Name, filePdf)

# createTitle('Titre', 'Nicolas K.', 'French')

from PyPDF2 import PdfFileMerger, PdfFileReader

def addPrefix(inputName, outputName):
	fileTitle = '/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/txt2pdf/Title_after.pdf'
	fileEmptyPage = '/Users/nicolas/Desktop/NaturaLingua/SubsToBilingualText/txt2pdf/empty.pdf'
	filenames=[fileTitle, fileEmptyPage,inputName]
	merger = PdfFileMerger()
	for filename in filenames:
		merger.append(PdfFileReader(open(filename, 'rb')))
	merger.write(outputName)

def getTitle(input):
	return input.split('/')[-1][:-4]

# input='/Users/nicolas/Desktop/NaturaLingua/pdfs_results/portuguese/movies/Carandiru.pdf'
def getOutput(input):
	newRoot='/Users/nicolas/Desktop/NaturaLingua/toEnglish'
	machin = input.split('/')
	result=f'{newRoot}/{machin[-3]}/{machin[-2]}/{machin[-1]}'
	return result

def createNecessaryFolders(output):
	machins = output.split('/')
	path=''
	for machin in machins[:-1]:
		path+=machin+'/'
		if not os.path.isdir(path):
			os.mkdir(path)

def createTitleAll(language, author, input):
	title=getTitle(input)
	createTitle(title, author, language)
	output=getOutput(input)
	createNecessaryFolders(output)
	addPrefix(input, output)

# createTitleAll("Portuguese", "Miguel Gonçalves", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/portuguese/movies/Carandiru.pdf')



# =========================
# movies












# ==========================

docsInfoFrench=[
	["François Truffaut", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/french/movies/Jules et Jim.pdf'],
	["François Truffaut", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/french/movies/Les 400 coups.pdf'],
	["Jacques Demy", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/french/movies/Les parapluies de Cherbourg.pdf'],
	["Jean-Pierre Jeunet", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/french/movies/Amelie.pdf'],
]


docsInfoItalian=[
	["Dino Risi", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/italian/movies/Il sorpasso.pdf'],
	["Federico Fellini", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/italian/movies/La dolce vita.pdf'],
	["Vittorio De Sica", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/italian/movies/Ladri di biciclette.pdf'],
	["Edoardo Ponti", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/italian/movies/The life ahead.pdf'],
]

docsInfoGreek=[
	["Peter Yates", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/grec/movies/Eleni.pdf'],
	["Theo Angelopoulos", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/grec/movies/Ulysses gaze.pdf'],
]

docsInfoTurkish=[
	["Ömer Faruk Sorak", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/turkish/movies/Gora.pdf'],
	["Yılmaz Erdoğan", "/Users/nicolas/Desktop/NaturaLingua/pdfs_results/turkish/movies/The butterfly's dream.pdf"],
	["Yılmaz Erdoğan", '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/turkish/movies/Vizontele.pdf'],
]

docsInfoHebrew=[
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E1.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E2.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E3.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E4.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E5.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E6.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E7.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E8.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E9.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E10.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E11.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S1E12.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E1.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E2.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E3.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E4.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E5.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E7.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E8.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E9.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E10.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E11.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S2E12.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S3E1.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S3E2.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S3E3.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S3E4.pdf'],
	['Ori Elon and Yehonatan Indursky', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hebrew/movies/Shtisel - S3E5.pdf'],
]

docsInfoArabic=[
	['Maroun Bagdadi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Beirut Oh Beirut (arabic script).pdf'],
	['Maroun Bagdadi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Beirut Oh Beirut (latin script).pdf'],
	['Maroun Bagdadi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Beirut Oh Beirut (arabic+latin script).pdf'],
	['Amin Dora', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Ghadi (arabic script).pdf'],
	['Amin Dora', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Ghadi (latin script).pdf'],
	['Amin Dora', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Ghadi (arabic+latin script).pdf'],
	['Lucien Bourjeily', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Heaven without people (arabic script).pdf'],
	['Lucien Bourjeily', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Heaven without people (latin script).pdf'],
	['Lucien Bourjeily', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Heaven without people (arabic+latin script).pdf'],
	['Randa Chahal Sabbag', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/The Kite (arabic script).pdf'],
	['Randa Chahal Sabbag', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/The Kite (latin script).pdf'],
	['Randa Chahal Sabbag', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/The Kite (arabic+latin script).pdf'],
	['Mir-Jean Bou Chaaya', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Very big shot (arabic script).pdf'],
	['Mir-Jean Bou Chaaya', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Very big shot (latin script).pdf'],
	['Mir-Jean Bou Chaaya', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/arabe/movies/Very big shot (arabic+latin script).pdf'],
]


docsInfoPersian=[
	['Asghar Farhadi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/persian/movies/A separation.pdf'],
	['Kamal Tabrizi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/persian/movies/Leili baa man ast.pdf'],
	['Bahman Ghobadi', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/persian/movies/Turtles can fly.pdf'],
]

docsInfoJapanese=[
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/The Howl Moving Castle.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Castle in the Sky.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Chihiro.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Kiki\'s Delivery Service.pdf'],
	['Gorō Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/From Up on Poppy Hill.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/The Wind Rises.pdf'],
	['Mamoru Hosoda', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Mirai no Mirai.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Mononoke.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Nausicaa.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Ponyo.pdf'],
	['Akira Kurosawa', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Rashomon.pdf'],
	['Satoshi Kon and Shôgo Furuya', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Tokyo godfathers.pdf'],
	['Hayao Miyazaki', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/japanese/movies/Totoro.pdf'],
]

docsInfoChinese=[
	['Edward Yang', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/chinese/movies/A brighter summer day.pdf'],
	['Xuan Liang and Chun Zhang', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/chinese/movies/Big fish and begonia.pdf'],
	['Jiang Wen', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/chinese/movies/In the heat of the sun.pdf'],
	['Rene Liu', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/chinese/movies/Us and them.pdf'],
]

docsInfoKorean=[
	['Kim Eun-sook and Kim Won-seok', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/korean/movies/Descendants of the sun - E1.pdf'],
	['Park Hoon-Jung', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/korean/movies/Night in paradise.pdf'],
	['Bong Joon-ho', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/korean/movies/Okja.pdf'],
]

docsInfoVietnamese=[
	['Lê Văn Kiệt', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/viet/movies/Furie.pdf'],
	['Tran Anh Hung', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/viet/movies/The Scent of Green Papaya.pdf'],
]

docsInfoHindi=[
	['Anurag Kashyap', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hindi/movies/Black Friday.pdf'],
	['Ruchi Narain', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hindi/movies/Guilty.pdf'],
	['Abhishek Chaubey', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/hindi/movies/Udta Punjab.pdf'],
]

docsInfoRussian=[
	['Kirill Serebrennikov', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/russe/movies/Leto.pdf'],
	['Andreï Tarkovski', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/russe/movies/Stalker.pdf'],
	['Mikhail Kalatozov', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/russe/movies/The cranes are flying.pdf'],
]

def doAllInList(language, docsInfos):
	for l in docsInfos:
		createTitleAll(language, l[0], l[1])

# todo -----> Scenario by : / Directed by :
# doAllInList('Russian', docsInfoRussian)

# ------------------------------
# youtube

def doSomething_2():
	root='/Users/nicolas/Desktop/NaturaLingua/pdfs_results'
	obj = os.scandir(root)
	for entry in obj :
		if entry.is_dir():
			folderName=f'{entry.path}/youtube'
			language=entry.name
			obj2 = os.scandir(folderName)
			for entry2 in obj2 :
				if entry2.is_file() and entry2.path.endswith('.pdf'):
					print(entry2.path)
					output=getOutput(entry2.path)
					# if(not os.path.isfile(output)):
					createTitleAll(language, '(YouTube channel)', entry2.path)

# createTitleAll('Korean', '(YouTube channel)', '/Users/nicolas/Desktop/NaturaLingua/pdfs_results/Korean/youtube/In a nutshell.pdf')

